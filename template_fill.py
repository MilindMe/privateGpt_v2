from docx import Document
from datetime import datetime

def fill_template(template_path, output_path, context): 
    doc = Document(template_path)
  

    for paragraph in doc.paragraphs: 
        for key, value in context.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    for table in doc.tables: 
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs: 
                    for key, value in context.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    doc.save(output_path)

if __name__ == "__main__":
    template_path = 'template.docx'
    output_path = 'filled_template.docx'

    context = {
        '{name}' : 'John Doe',
        '{address}' : '1234 Main St, Somewhere, Mauritius',
        '{date}': datetime.today().strftime('%B %d, %Y')
    } 

    fill_template(template_path, output_path, context)
    print(f'The document has been filled and saved to {output_path}')
